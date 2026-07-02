#!/usr/bin/env python3
"""OmniPMX document/runtime MCP server.

Provides xskills-compatible MCP tools for:
- OCR image URL/base64
- OCR PDF URL/base64
- PDF text extraction with OCR fallback
- Markdown/text to DOCX
- Markdown/text to PDF
- CSV/table to PNG chart
- Optional restricted Python execution

Endpoint:
    POST /mcp

Important environment variables:
    MCP_API_TOKEN             Optional bearer token.
    ENABLE_CODE_EXECUTION     Set to true to enable python_execute.
    CODE_EXECUTION_TIMEOUT    Seconds, default 15.
"""

from __future__ import annotations

import base64
import csv
import io
import json
import os
import re
import subprocess
import tempfile
import textwrap
import time
import traceback
import urllib.request
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from typing import Any


SERVER_NAME = "omnipmx-document-runtime-mcp"
SERVER_VERSION = "1.0.20260702"
PROTOCOL_VERSION = "2025-06-18"
DEFAULT_TIMEOUT = 40
MAX_FETCH_BYTES = 40_000_000


def env(name: str, default: str = "") -> str:
    value = os.environ.get(name, "").strip()
    return value if value else default


def env_int(name: str, default: int) -> int:
    try:
        return int(env(name, str(default)))
    except ValueError:
        return default


def dumps(data: Any) -> bytes:
    return json.dumps(data, ensure_ascii=False, separators=(",", ":")).encode("utf-8")


def rpc_result(request_id: Any, result: Any) -> dict[str, Any]:
    return {"jsonrpc": "2.0", "id": request_id, "result": result}


def rpc_error(request_id: Any, code: int, message: str, data: Any | None = None) -> dict[str, Any]:
    error: dict[str, Any] = {"code": code, "message": message}
    if data is not None:
        error["data"] = data
    return {"jsonrpc": "2.0", "id": request_id, "error": error}


def get_url(url: str, max_bytes: int = MAX_FETCH_BYTES) -> tuple[bytes, str, str]:
    if not re.match(r"^https?://", url, flags=re.I):
        raise ValueError("url must start with http:// or https://")
    req = urllib.request.Request(
        url,
        headers={
            "User-Agent": f"{SERVER_NAME}/{SERVER_VERSION}",
            "Accept": "application/pdf,image/*,text/html,text/plain,*/*",
        },
        method="GET",
    )
    with urllib.request.urlopen(req, timeout=DEFAULT_TIMEOUT) as resp:
        raw = resp.read(max_bytes + 1)
        content_type = resp.headers.get("content-type", "")
        final_url = resp.geturl()
    return raw[:max_bytes], content_type, final_url


def bytes_from_args(args: dict[str, Any], *, url_key: str = "url", b64_key: str = "base64") -> tuple[bytes, str, str]:
    if args.get(b64_key):
        return base64.b64decode(str(args[b64_key])), str(args.get("content_type", "")), "base64-input"
    if args.get(url_key):
        return get_url(str(args[url_key]).strip())
    raise ValueError(f"provide {url_key} or {b64_key}")


def out_file(filename: str, mime_type: str, raw: bytes, extra: dict[str, Any] | None = None) -> dict[str, Any]:
    data = {
        "filename": filename,
        "mime_type": mime_type,
        "byte_count": len(raw),
        "base64": base64.b64encode(raw).decode("ascii"),
        "retrieved_at_epoch": int(time.time()),
    }
    if extra:
        data.update(extra)
    return data


def ocr_image(args: dict[str, Any]) -> dict[str, Any]:
    from PIL import Image
    import pytesseract

    raw, content_type, source = bytes_from_args(args)
    lang = str(args.get("lang", "eng+chi_sim")).strip() or "eng"
    image = Image.open(io.BytesIO(raw))
    text = pytesseract.image_to_string(image, lang=lang)
    return {
        "status": "OCR_DONE",
        "source": source,
        "content_type": content_type,
        "lang": lang,
        "text": text,
        "char_count": len(text),
    }


def pdf_text_extract(args: dict[str, Any]) -> dict[str, Any]:
    from pypdf import PdfReader

    raw, content_type, source = bytes_from_args(args)
    max_pages = max(1, min(int(args.get("max_pages", 200)), 1000))
    max_chars = max(1000, min(int(args.get("max_chars", 300000)), 1_000_000))
    reader = PdfReader(io.BytesIO(raw))
    page_texts = []
    for idx, page in enumerate(reader.pages[:max_pages], start=1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:  # noqa: BLE001
            text = f"[PAGE_EXTRACTION_ERROR: {exc}]"
        if text.strip():
            page_texts.append({"page": idx, "text": text.strip()})
    joined = "\n\n".join(f"--- page {p['page']} ---\n{p['text']}" for p in page_texts)
    truncated = len(joined) > max_chars or len(reader.pages) > max_pages
    return {
        "status": "PDF_TEXT_EXTRACTED" if page_texts else "NO_EXTRACTABLE_TEXT",
        "source": source,
        "content_type": content_type,
        "page_count": len(reader.pages),
        "pages_attempted": min(max_pages, len(reader.pages)),
        "pages_with_text": len(page_texts),
        "text": joined[:max_chars],
        "truncated": truncated,
    }


def ocr_pdf(args: dict[str, Any]) -> dict[str, Any]:
    import fitz
    from PIL import Image
    import pytesseract

    raw, content_type, source = bytes_from_args(args)
    lang = str(args.get("lang", "eng+chi_sim")).strip() or "eng"
    dpi = max(100, min(int(args.get("dpi", 180)), 300))
    max_pages = max(1, min(int(args.get("max_pages", 60)), 300))
    max_chars = max(1000, min(int(args.get("max_chars", 300000)), 1_000_000))
    doc = fitz.open(stream=raw, filetype="pdf")
    page_texts = []
    zoom = dpi / 72.0
    matrix = fitz.Matrix(zoom, zoom)
    for idx in range(min(max_pages, doc.page_count)):
        page = doc.load_page(idx)
        pix = page.get_pixmap(matrix=matrix, alpha=False)
        image = Image.open(io.BytesIO(pix.tobytes("png")))
        text = pytesseract.image_to_string(image, lang=lang)
        if text.strip():
            page_texts.append({"page": idx + 1, "text": text.strip()})
    joined = "\n\n".join(f"--- page {p['page']} ---\n{p['text']}" for p in page_texts)
    return {
        "status": "PDF_OCR_DONE" if page_texts else "NO_OCR_TEXT",
        "source": source,
        "content_type": content_type,
        "page_count": doc.page_count,
        "pages_attempted": min(max_pages, doc.page_count),
        "pages_with_text": len(page_texts),
        "lang": lang,
        "dpi": dpi,
        "text": joined[:max_chars],
        "truncated": len(joined) > max_chars or doc.page_count > max_pages,
    }


def pdf_extract_or_ocr(args: dict[str, Any]) -> dict[str, Any]:
    result = pdf_text_extract(args)
    min_chars = int(args.get("ocr_if_text_below_chars", 200))
    if result.get("status") == "NO_EXTRACTABLE_TEXT" or len(str(result.get("text", ""))) < min_chars:
        ocr_result = ocr_pdf(args)
        ocr_result["text_extract_status_before_ocr"] = result.get("status")
        return ocr_result
    return result


def split_markdown_lines(markdown: str) -> list[tuple[str, str]]:
    lines: list[tuple[str, str]] = []
    for raw in markdown.splitlines():
        line = raw.rstrip()
        if line.startswith("# "):
            lines.append(("h1", line[2:].strip()))
        elif line.startswith("## "):
            lines.append(("h2", line[3:].strip()))
        elif line.startswith("### "):
            lines.append(("h3", line[4:].strip()))
        elif line.startswith("- ") or line.startswith("* "):
            lines.append(("bullet", line[2:].strip()))
        elif re.match(r"^\d+\.\s+", line):
            lines.append(("number", re.sub(r"^\d+\.\s+", "", line).strip()))
        elif not line.strip():
            lines.append(("blank", ""))
        else:
            lines.append(("p", line.strip()))
    return lines


def markdown_to_docx(args: dict[str, Any]) -> dict[str, Any]:
    from docx import Document
    from docx.shared import Pt

    markdown = str(args.get("markdown", "") or args.get("text", "")).strip()
    if not markdown:
        raise ValueError("markdown or text is required")
    filename = str(args.get("filename", "omnipmx_report_1.0_20260702.docx"))
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    for kind, text in split_markdown_lines(markdown):
        if kind == "h1":
            doc.add_heading(text, level=1)
        elif kind == "h2":
            doc.add_heading(text, level=2)
        elif kind == "h3":
            doc.add_heading(text, level=3)
        elif kind == "bullet":
            doc.add_paragraph(text, style="List Bullet")
        elif kind == "number":
            doc.add_paragraph(text, style="List Number")
        elif kind == "blank":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return out_file(filename, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", buf.getvalue())


def markdown_to_pdf(args: dict[str, Any]) -> dict[str, Any]:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont

    markdown = str(args.get("markdown", "") or args.get("text", "")).strip()
    if not markdown:
        raise ValueError("markdown or text is required")
    filename = str(args.get("filename", "omnipmx_report_1.0_20260702.pdf"))
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    base = ParagraphStyle("OmniBase", parent=styles["BodyText"], fontName="STSong-Light", fontSize=10, leading=14)
    h1 = ParagraphStyle("OmniH1", parent=base, fontSize=18, leading=24, spaceAfter=12)
    h2 = ParagraphStyle("OmniH2", parent=base, fontSize=14, leading=18, spaceBefore=10, spaceAfter=8)
    h3 = ParagraphStyle("OmniH3", parent=base, fontSize=12, leading=16, spaceBefore=8, spaceAfter=6)
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=1.8 * cm, rightMargin=1.8 * cm, topMargin=1.6 * cm, bottomMargin=1.6 * cm)
    story: list[Any] = []
    bullet_items: list[ListItem] = []

    def flush_bullets() -> None:
        nonlocal bullet_items
        if bullet_items:
            story.append(ListFlowable(bullet_items, bulletType="bullet"))
            bullet_items = []

    for kind, text in split_markdown_lines(markdown):
        safe = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if kind == "bullet":
            bullet_items.append(ListItem(Paragraph(safe, base)))
            continue
        flush_bullets()
        if kind == "h1":
            story.append(Paragraph(safe, h1))
        elif kind == "h2":
            story.append(Paragraph(safe, h2))
        elif kind == "h3":
            story.append(Paragraph(safe, h3))
        elif kind == "blank":
            story.append(Spacer(1, 0.2 * cm))
        else:
            story.append(Paragraph(safe, base))
    flush_bullets()
    doc.build(story)
    return out_file(filename, "application/pdf", buf.getvalue())


def csv_to_plot_png(args: dict[str, Any]) -> dict[str, Any]:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    csv_text = str(args.get("csv", "")).strip()
    if not csv_text:
        raise ValueError("csv is required")
    x_field = str(args.get("x_field", "")).strip()
    y_field = str(args.get("y_field", "")).strip()
    chart_type = str(args.get("chart_type", "bar")).strip()
    filename = str(args.get("filename", "omnipmx_chart_1.0_20260702.png"))
    rows = list(csv.DictReader(io.StringIO(csv_text)))
    if not rows:
        raise ValueError("csv has no rows")
    if not x_field:
        x_field = list(rows[0].keys())[0]
    if not y_field:
        y_field = list(rows[0].keys())[1]
    x = [r.get(x_field, "") for r in rows]
    y = [float(r.get(y_field, "0") or 0) for r in rows]
    fig, ax = plt.subplots(figsize=(8, 4.5), dpi=160)
    if chart_type == "line":
        ax.plot(x, y, marker="o")
    else:
        ax.bar(x, y)
    ax.set_xlabel(x_field)
    ax.set_ylabel(y_field)
    ax.set_title(str(args.get("title", "")))
    ax.tick_params(axis="x", rotation=45)
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png")
    plt.close(fig)
    return out_file(filename, "image/png", buf.getvalue())


def python_execute(args: dict[str, Any]) -> dict[str, Any]:
    if env("ENABLE_CODE_EXECUTION", "false").lower() not in {"1", "true", "yes"}:
        return {
            "status": "CODE_EXECUTION_DISABLED",
            "message": "Set ENABLE_CODE_EXECUTION=true on the MCP server to enable this tool.",
        }
    code = str(args.get("code", ""))
    if not code.strip():
        raise ValueError("code is required")
    timeout = max(1, min(int(args.get("timeout", env_int("CODE_EXECUTION_TIMEOUT", 15))), 60))
    with tempfile.TemporaryDirectory() as tmp:
        script = os.path.join(tmp, "script.py")
        with open(script, "w", encoding="utf-8") as f:
            f.write(code)
        proc = subprocess.run(
            ["python", script],
            cwd=tmp,
            text=True,
            capture_output=True,
            timeout=timeout,
            check=False,
        )
    return {
        "status": "EXECUTED",
        "returncode": proc.returncode,
        "stdout": proc.stdout[-20000:],
        "stderr": proc.stderr[-20000:],
        "timeout_seconds": timeout,
    }


TOOLS: dict[str, dict[str, Any]] = {
    "ocr_image": {
        "description": "OCR an image from URL or base64. Supports English and simplified Chinese if installed in Docker.",
        "inputSchema": {
            "type": "object",
            "properties": {"url": {"type": "string"}, "base64": {"type": "string"}, "content_type": {"type": "string"}, "lang": {"type": "string", "default": "eng+chi_sim"}},
        },
        "handler": ocr_image,
    },
    "pdf_text_extract": {
        "description": "Extract embedded text from a PDF URL/base64 using pypdf. Not OCR.",
        "inputSchema": {
            "type": "object",
            "properties": {"url": {"type": "string"}, "base64": {"type": "string"}, "max_pages": {"type": "integer", "default": 200}, "max_chars": {"type": "integer", "default": 300000}},
        },
        "handler": pdf_text_extract,
    },
    "ocr_pdf": {
        "description": "OCR PDF pages from URL/base64 using PyMuPDF rendering and Tesseract.",
        "inputSchema": {
            "type": "object",
            "properties": {"url": {"type": "string"}, "base64": {"type": "string"}, "lang": {"type": "string", "default": "eng+chi_sim"}, "dpi": {"type": "integer", "default": 180}, "max_pages": {"type": "integer", "default": 60}, "max_chars": {"type": "integer", "default": 300000}},
        },
        "handler": ocr_pdf,
    },
    "pdf_extract_or_ocr": {
        "description": "Extract PDF embedded text, and OCR automatically if text is absent/too short.",
        "inputSchema": {
            "type": "object",
            "properties": {"url": {"type": "string"}, "base64": {"type": "string"}, "lang": {"type": "string", "default": "eng+chi_sim"}, "max_pages": {"type": "integer", "default": 60}, "max_chars": {"type": "integer", "default": 300000}, "ocr_if_text_below_chars": {"type": "integer", "default": 200}},
        },
        "handler": pdf_extract_or_ocr,
    },
    "markdown_to_docx": {
        "description": "Generate a DOCX file from Markdown/text and return base64 file content.",
        "inputSchema": {"type": "object", "properties": {"markdown": {"type": "string"}, "filename": {"type": "string"}}, "required": ["markdown"]},
        "handler": markdown_to_docx,
    },
    "markdown_to_pdf": {
        "description": "Generate a PDF file from Markdown/text and return base64 file content.",
        "inputSchema": {"type": "object", "properties": {"markdown": {"type": "string"}, "filename": {"type": "string"}}, "required": ["markdown"]},
        "handler": markdown_to_pdf,
    },
    "csv_to_plot_png": {
        "description": "Render a simple bar/line chart from CSV text and return base64 PNG.",
        "inputSchema": {"type": "object", "properties": {"csv": {"type": "string"}, "x_field": {"type": "string"}, "y_field": {"type": "string"}, "chart_type": {"type": "string", "default": "bar"}, "title": {"type": "string"}, "filename": {"type": "string"}}, "required": ["csv"]},
        "handler": csv_to_plot_png,
    },
    "python_execute": {
        "description": "Optional restricted Python execution. Disabled unless ENABLE_CODE_EXECUTION=true.",
        "inputSchema": {"type": "object", "properties": {"code": {"type": "string"}, "timeout": {"type": "integer", "default": 15}}, "required": ["code"]},
        "handler": python_execute,
    },
}


def list_tools() -> list[dict[str, Any]]:
    return [{"name": n, "description": s["description"], "inputSchema": s["inputSchema"]} for n, s in TOOLS.items()]


def call_tool(name: str, arguments: dict[str, Any]) -> dict[str, Any]:
    if name not in TOOLS:
        raise ValueError(f"unknown tool: {name}")
    data = TOOLS[name]["handler"](arguments or {})
    return {"content": [{"type": "text", "text": json.dumps(data, ensure_ascii=False, indent=2)}], "isError": False}


def handle_rpc(message: dict[str, Any]) -> dict[str, Any] | None:
    request_id = message.get("id")
    method = message.get("method")
    params = message.get("params") or {}
    try:
        if method == "initialize":
            return rpc_result(
                request_id,
                {
                    "protocolVersion": PROTOCOL_VERSION,
                    "capabilities": {"tools": {"listChanged": False}},
                    "serverInfo": {"name": SERVER_NAME, "version": SERVER_VERSION},
                },
            )
        if method == "notifications/initialized":
            return None
        if method == "tools/list":
            return rpc_result(request_id, {"tools": list_tools()})
        if method == "tools/call":
            return rpc_result(request_id, call_tool(str(params.get("name", "")), params.get("arguments") or {}))
        if method == "ping":
            return rpc_result(request_id, {})
        return rpc_error(request_id, -32601, f"method not found: {method}")
    except subprocess.TimeoutExpired:
        return rpc_error(request_id, -32000, "execution timed out")
    except Exception as exc:  # noqa: BLE001
        return rpc_error(request_id, -32000, str(exc), traceback.format_exc())


class Handler(BaseHTTPRequestHandler):
    server_version = SERVER_NAME

    def do_HEAD(self) -> None:  # noqa: N802
        if self.path in {"/", "/health", "/mcp"}:
            self.send_response(200)
            self.send_common_headers()
            self.end_headers()
            return
        self.send_response(404)
        self.send_common_headers()
        self.end_headers()

    def do_OPTIONS(self) -> None:  # noqa: N802
        self.send_response(204)
        self.send_common_headers()
        self.end_headers()

    def do_GET(self) -> None:  # noqa: N802
        if self.path in {"/", "/health"}:
            self.send_json(200, {"status": "ok", "server": SERVER_NAME, "version": SERVER_VERSION})
        elif self.path == "/mcp":
            self.send_json(200, {"status": "ok", "message": "Use POST /mcp with JSON-RPC.", "server": SERVER_NAME, "version": SERVER_VERSION, "tools": list(TOOLS.keys())})
        else:
            self.send_json(405, {"error": "Use POST /mcp for MCP JSON-RPC.", "health": "/health"})

    def do_POST(self) -> None:  # noqa: N802
        if self.path != "/mcp":
            self.send_json(404, {"error": "not found", "mcp_endpoint": "/mcp"})
            return
        if not self.authorized():
            self.send_json(401, {"error": "unauthorized"})
            return
        try:
            length = int(self.headers.get("content-length", "0"))
            payload = json.loads(self.rfile.read(length).decode("utf-8"))
        except Exception as exc:  # noqa: BLE001
            self.send_json(400, rpc_error(None, -32700, f"invalid JSON: {exc}"))
            return
        if isinstance(payload, list):
            responses = [resp for item in payload if (resp := handle_rpc(item)) is not None]
            self.send_json(200 if responses else 202, responses if responses else {})
            return
        response = handle_rpc(payload)
        if response is None:
            self.send_response(202)
            self.send_common_headers()
            self.end_headers()
            return
        self.send_json(200, response)

    def authorized(self) -> bool:
        token = env("MCP_API_TOKEN")
        if not token:
            return True
        return self.headers.get("authorization", "") == f"Bearer {token}"

    def send_json(self, status: int, data: Any) -> None:
        body = dumps(data)
        self.send_response(status)
        self.send_common_headers()
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def send_common_headers(self) -> None:
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "GET, POST, HEAD, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type, Accept, Authorization, Mcp-Session-Id")
        self.send_header("Access-Control-Expose-Headers", "Mcp-Session-Id")

    def log_message(self, fmt: str, *args: Any) -> None:
        if env("MCP_QUIET_LOGS", "0") == "1":
            return
        super().log_message(fmt, *args)


def main() -> None:
    host = env("HOST", "0.0.0.0")
    port = env_int("PORT", 8790)
    server = ThreadingHTTPServer((host, port), Handler)
    print(f"{SERVER_NAME} listening on http://{host}:{port}/mcp", flush=True)
    server.serve_forever()


if __name__ == "__main__":
    main()
